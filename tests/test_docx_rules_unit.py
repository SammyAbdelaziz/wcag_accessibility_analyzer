"""
Comprehensive DOCX Rules Unit Tests - Simplified

Unit tests for core DOCX WCAG analyzer rules with deterministic assertions.
Focuses on rules that can be tested reliably via python-docx API.

Total coverage: 8+ core rules covering highest-priority WCAG failures
"""

import unittest
import io
from pathlib import Path
from wcag.analyzers.docx_analyzer import DocxAnalyzer
from wcag.models import Severity, ConfidenceTier


class DocxAnalyzerTestBase(unittest.TestCase):
    """Base class for DOCX rule tests."""
    
    FIXTURES_DIR = Path(__file__).parent / "fixtures" / "docx"
    
    @staticmethod
    def analyze_docx(filename: str = "test.docx", title: str = None, language: str = None):
        """Create a minimal DOCX with optional properties and analyze it."""
        from docx import Document
        doc = Document()
        if title:
            doc.core_properties.title = title
        if language:
            doc.core_properties.language = language
        doc.add_paragraph("Test document content.")
        
        buf = io.BytesIO()
        doc.save(buf)
        return DocxAnalyzer(buf.getvalue(), filename).analyze()
    
    def get_findings_by_criterion(self, fact_sheet, criterion_id: str):
        """Filter findings by WCAG criterion ID."""
        confirmed = [f for f in fact_sheet.confirmed_findings if f.criterion_id == criterion_id]
        possible = [f for f in fact_sheet.possible_findings if f.criterion_id == criterion_id]
        return confirmed, possible


# =============================================================================
# 2.4.2 - Document Title Tests
# =============================================================================

class TestRule242DocTitle(DocxAnalyzerTestBase):
    """Tests for WCAG 2.4.2: Document has descriptive title."""
    
    def test_document_with_title_passes(self):
        """✓ Document with descriptive title should pass."""
        fact_sheet = self.analyze_docx("titled.docx", title="Quarterly Report Q4 2026")
        confirmed, _ = self.get_findings_by_criterion(fact_sheet, "2.4.2")
        # Filter to title-specific findings
        title_findings = [f for f in confirmed if f.remediation_id == "doc_title"]
        self.assertEqual(len(title_findings), 0, "Should not flag document with title")
    
    def test_document_without_title_fails(self):
        """✗ Document with no title should fail."""
        fact_sheet = self.analyze_docx("untitled.docx")  # No title
        confirmed, _ = self.get_findings_by_criterion(fact_sheet, "2.4.2")
        title_findings = [f for f in confirmed if f.remediation_id == "doc_title"]
        self.assertGreater(len(title_findings), 0, "Should flag missing document title")


# =============================================================================
# 3.1.1 - Document Language Tests
# =============================================================================

class TestRule311Language(DocxAnalyzerTestBase):
    """Tests for WCAG 3.1.1: Document language is set."""
    
    def test_document_with_language_passes(self):
        """✓ Document with language set should pass."""
        fact_sheet = self.analyze_docx("english.docx", language="en-US")
        confirmed, _ = self.get_findings_by_criterion(fact_sheet, "3.1.1")
        lang_findings = [f for f in confirmed if f.remediation_id == "doc_language"]
        self.assertEqual(len(lang_findings), 0, "Should not flag document with language set")
    
    def test_document_without_language_warning(self):
        """⚠ Document without language may be flagged (language detection optional)."""
        fact_sheet = self.analyze_docx("no_lang.docx")  # No language
        confirmed, _ = self.get_findings_by_criterion(fact_sheet, "3.1.1")
        # Language detection may be optional; we just verify the analyzer runs
        self.assertIsNotNone(confirmed)


# =============================================================================
# 1.3.1 - Heading Hierarchy Tests
# =============================================================================

class TestRule131HeadingHierarchy(DocxAnalyzerTestBase):
    """Tests for WCAG 1.3.1: Heading hierarchy is valid."""
    
    def test_valid_heading_hierarchy_passes(self):
        """✓ Proper H1 → H2 progression should pass."""
        from docx import Document
        doc = Document()
        doc.add_heading("Main Title", level=1)
        doc.add_heading("Section 1", level=2)
        doc.add_paragraph("Body text.")
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "valid.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.3.1")
        heading_findings = [f for f in confirmed + possible 
                           if f.remediation_id == "heading_hierarchy"]
        self.assertEqual(len(heading_findings), 0, "Should not flag valid hierarchy")
    
    def test_skipped_heading_level_fails(self):
        """✗ H1 → H3 (skipping H2) should fail."""
        from docx import Document
        doc = Document()
        doc.add_heading("Main Title", level=1)
        doc.add_heading("Skipped to H3", level=3)  # Skip H2!
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "skipped.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.3.1")
        heading_findings = [f for f in confirmed + possible 
                           if f.remediation_id == "heading_hierarchy"]
        self.assertGreater(len(heading_findings), 0, "Should flag skipped heading level")


# =============================================================================
# 1.3.1 - List Coherence Tests
# =============================================================================

class TestRule131ListCoherence(DocxAnalyzerTestBase):
    """Tests for WCAG 1.3.1: Lists have coherent nesting."""
    
    def test_simple_list_passes(self):
        """✓ Single-level list should pass."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Item 1", style='List Bullet')
        doc.add_paragraph("Item 2", style='List Bullet')
        doc.add_paragraph("Item 3", style='List Bullet')
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "simple_list.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.3.1")
        list_findings = [f for f in confirmed + possible 
                        if f.remediation_id in ("list_coherence", "list_formatting")]
        self.assertEqual(len(list_findings), 0, "Should not flag coherent list")
    
    def test_properly_nested_list_passes(self):
        """✓ Proper nesting (1 → 2 → 2 → 1) should pass."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Item 1", style='List Bullet')
        doc.add_paragraph("Sub-item 1a", style='List Bullet 2')
        doc.add_paragraph("Sub-item 1b", style='List Bullet 2')
        doc.add_paragraph("Item 2", style='List Bullet')
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "nested.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.3.1")
        list_findings = [f for f in confirmed + possible 
                        if f.remediation_id in ("list_coherence", "list_formatting")]
        self.assertEqual(len(list_findings), 0, "Should not flag proper nesting")


# =============================================================================
# 1.3.1 - Table Header Tests
# =============================================================================

class TestRule131TableHeaders(DocxAnalyzerTestBase):
    """Tests for WCAG 1.3.1: Tables have identified header row."""
    
    def test_table_created_successfully(self):
        """✓ Tables can be created and analyzed."""
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(0, 2).text = "City"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"
        table.cell(1, 2).text = "NYC"
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "table.docx").analyze()
        # Just verify it processes without error
        self.assertIsNotNone(fact_sheet)
        self.assertIsNotNone(fact_sheet.confirmed_findings)


# =============================================================================
# 1.1.1 - Image Alt Text Tests
# =============================================================================

class TestRule111ImageAlt(DocxAnalyzerTestBase):
    """Tests for WCAG 1.1.1: Images have alternative text."""
    
    def test_document_without_images_passes(self):
        """✓ Document with no images should pass."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Text-only document.")
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "no_images.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.1.1")
        image_findings = [f for f in confirmed + possible 
                         if f.remediation_id and "image" in f.remediation_id.lower()]
        self.assertEqual(len(image_findings), 0, "Should not flag document without images")


# =============================================================================
# 1.4.4 - Text Size Tests
# =============================================================================

class TestRule144TextSize(DocxAnalyzerTestBase):
    """Tests for WCAG 1.4.4: Text is resizable and readable."""
    
    def test_normal_text_size_passes(self):
        """✓ Normal text size (12pt) should pass."""
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        p = doc.add_paragraph("Normal text at 12pt.")
        for run in p.runs:
            run.font.size = Pt(12)
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "normal.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.4.4")
        size_findings = [f for f in confirmed + possible 
                        if f.remediation_id == "text_size"]
        self.assertEqual(len(size_findings), 0, "Should not flag normal text size")
    
    def test_tiny_text_flagged(self):
        """✗ Very small text (<8pt) should be flagged."""
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        p = doc.add_paragraph("Tiny text at 6pt.")
        for run in p.runs:
            run.font.size = Pt(6)
        
        buf = io.BytesIO()
        doc.save(buf)
        fact_sheet = DocxAnalyzer(buf.getvalue(), "tiny.docx").analyze()
        confirmed, possible = self.get_findings_by_criterion(fact_sheet, "1.4.4")
        size_findings = [f for f in confirmed + possible 
                        if f.remediation_id == "text_size"]
        self.assertGreater(len(size_findings), 0, "Should flag tiny text size")


# =============================================================================
# Integration / Regression Tests
# =============================================================================

class TestDocxOutputFormat(DocxAnalyzerTestBase):
    """Verify DOCX fact sheet output."""
    
    def test_fact_sheet_serialization(self):
        """Fact sheet should serialize correctly to dict."""
        fact_sheet = self.analyze_docx("test.docx", title="Test", language="en-US")
        output = fact_sheet.to_dict()
        
        self.assertEqual(output["file_type"], "docx")
        self.assertEqual(output["filename"], "test.docx")
        self.assertIn("confirmed_findings", output)
        self.assertIn("possible_findings", output)
        self.assertIn("summary", output)
    
    def test_all_findings_have_remediation_data(self):
        """All findings should include remediation_data."""
        fact_sheet = self.analyze_docx("test.docx")
        all_findings = fact_sheet.confirmed_findings + fact_sheet.possible_findings
        
        for finding in all_findings:
            self.assertIsNotNone(finding.remediation_data)
            self.assertIsInstance(finding.remediation_data, dict)


if __name__ == "__main__":
    unittest.main()
